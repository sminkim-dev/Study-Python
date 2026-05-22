from docx import Document

doc = Document()

# =====================
# 표지
# =====================
doc.add_heading('컴공 자격증 로드맵 정리', 0)
doc.add_paragraph('대화 기반 정리본')
doc.add_paragraph('')
doc.add_paragraph('요약: 컴퓨터공학 전공자의 자격증 로드맵 (기본 / 취업 / 실무 단계)')
doc.add_page_break()

# =====================
# 목차
# =====================
doc.add_heading('목차', level=1)

toc = [
"1. 전체 개요",
"2. 자격증 전체 구조",
"3. 졸업 전 필수 라인",
"4. 취업 경쟁력 자격증",
"5. 실무/상위권 자격증",
"6. 추천 로드맵",
"7. 난이도 체감",
"8. 핵심 결론"
]

for item in toc:
    doc.add_paragraph(item)

doc.add_page_break()

# =====================
# 본문
# =====================
doc.add_heading('1. 전체 개요', level=1)
doc.add_paragraph("자격증은 기본 / 취업 / 실무 3단계로 나뉜다.")
doc.add_paragraph("핵심: 자격증은 필터 역할이며 실제 취업은 프로젝트가 중요하다.")

# =====================
# 표
# =====================
doc.add_heading('2. 자격증 전체 구조', level=1)

table = doc.add_table(rows=1, cols=5)
hdr = table.rows[0].cells
hdr[0].text = '구분'
hdr[1].text = '자격증'
hdr[2].text = '난이도'
hdr[3].text = '준비기간'
hdr[4].text = '취업 영향'

rows = [
("기본 필수","정보처리기사","중","1~3개월","매우 높음"),
("기본 필수","SQLD","하~중","2~3주","높음"),
("기본 필수","리눅스마스터 2급","하","2~4주","중"),
("취업 부스터","AWS CCP","중","2~4주","높음"),
("실무형","AWS SAA","중~상","1~3개월","매우 높음"),
("데이터","ADsP","하","2~3주","중"),
("상위권","정보보안기사","상","3~6개월","매우 높음"),
]

for r in rows:
    row_cells = table.add_row().cells
    for i, v in enumerate(r):
        row_cells[i].text = v

# =====================
# 나머지 섹션
# =====================
doc.add_heading('3. 졸업 전 필수 라인', level=1)
doc.add_paragraph("정보처리기사 / SQLD / 리눅스마스터 2급")

doc.add_heading('4. 취업 경쟁력 자격증', level=1)
doc.add_paragraph("네트워크관리사 2급 / AWS CCP / ADsP")

doc.add_heading('5. 실무/상위권 자격증', level=1)
doc.add_paragraph("AWS SAA / 정보보안기사")

doc.add_heading('6. 추천 로드맵', level=1)
doc.add_paragraph("1~2학년: SQLD, 리눅스마스터 2급, 네트워크관리사 2급")
doc.add_paragraph("2~3학년: 정보처리기사, AWS CCP")
doc.add_paragraph("3~4학년: AWS SAA, 보안/데이터 심화")

doc.add_heading('7. 난이도 체감', level=1)
doc.add_paragraph("하: 2~3주 / 중: 1~2달 / 상: 3개월 이상")

doc.add_heading('8. 핵심 결론', level=1)
doc.add_paragraph("자격증은 합격 필터이며, 실제 취업은 프로젝트가 핵심이다.")
doc.add_paragraph("최소 3종: 정보처리기사, SQLD, 리눅스마스터 2급")

# 저장
doc.save("컴공_자격증_로드맵.docx")

print("완료: 컴공_자격증_로드맵.docx 생성됨")